"""
Legal Document Placeholder Filler Web Application.

This Flask application allows users to upload legal documents,
identifies placeholders, and provides a conversational interface
to fill them in.
"""

import os
import re
import json
from datetime import datetime
from werkzeug.utils import secure_filename
from flask import Flask, render_template, request, jsonify, send_file, session
from docx import Document
from docx.shared import RGBColor


app = Flask(__name__)
app.secret_key = os.urandom(24)
app.config["UPLOAD_FOLDER"] = "/tmp/uploads"
app.config["OUTPUT_FOLDER"] = "/tmp/outputs"
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16MB max file size

# Ensure directories exist
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["OUTPUT_FOLDER"], exist_ok=True)

ALLOWED_EXTENSIONS = {"docx"}


def allowed_file(filename):
    """Check if file extension is allowed."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_placeholders(doc_path):
    """
    Extract placeholders from a Word document.

    Identifies patterns like [text], $[___], etc.
    Returns a list of unique placeholders with their context.
    Handles duplicate placeholders by tracking each occurrence.
    """
    doc = Document(doc_path)
    placeholders = []
    seen_with_context = {}  # Track placeholders with their context

    # Pattern to match various placeholder formats
    patterns = [
        r"\[([^\]]+)\]",  # [text]
        r"\$\[([^\]]+)\]",  # $[text]
    ]

    for para in doc.paragraphs:
        text = para.text
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                full_match = match.group(0)
                inner_text = match.group(1)

                # Clean up the inner text
                cleaned_name = inner_text.strip("_").strip()

                # Determine if it's currency
                is_currency = full_match.startswith("$")

                # Skip text placeholders that are empty after cleaning
                # But keep currency placeholders even if they're just underscores
                if not is_currency and cleaned_name == "":
                    continue

                # For underscore-only currency placeholders, give them a base name
                if is_currency and cleaned_name == "":
                    cleaned_name = "Amount"

                # Get context for this placeholder
                start = max(0, match.start() - 50)
                end = min(len(text), match.end() + 50)
                context = text[start:end].strip()

                # For currency placeholders without descriptive names, try to extract from context
                if is_currency and cleaned_name == "Amount":
                    context_lower = context.lower()
                    if "purchase amount" in context_lower:
                        cleaned_name = "Purchase Amount"
                    elif (
                        "valuation cap" in context_lower
                        or "post-money" in context_lower
                    ):
                        cleaned_name = "Post-Money Valuation Cap"
                    elif "investment" in context_lower:
                        cleaned_name = "Investment Amount"

                # Check if we've seen this exact placeholder before
                if full_match in seen_with_context:
                    # We have a duplicate - differentiate by context
                    seen_with_context[full_match].append(context)
                    occurrence_num = len(seen_with_context[full_match])

                    # If we haven't already named it, try to extract from context
                    if cleaned_name == "Amount":
                        context_lower = context.lower()
                        if "purchase amount" in context_lower:
                            display_name = "Purchase Amount"
                        elif (
                            "valuation cap" in context_lower
                            or "post-money" in context_lower
                        ):
                            display_name = "Post-Money Valuation Cap"
                        else:
                            display_name = f"{cleaned_name} (#{occurrence_num})"
                    else:
                        display_name = cleaned_name
                else:
                    # First occurrence
                    seen_with_context[full_match] = [context]
                    display_name = cleaned_name

                # Create placeholder info
                placeholder_info = {
                    "original": full_match,
                    "name": display_name,
                    "type": "currency" if is_currency else "text",
                    "context": context,
                }

                placeholders.append(placeholder_info)

    return placeholders


def get_placeholder_question(placeholder):
    """
    Generate a conversational question for a placeholder.
    """
    name = placeholder["name"]

    # Custom questions for common placeholders
    questions = {
        "Company Name": "What's the name of the company?",
        "Investor Name": "Who is the investor?",
        "Date of Safe": "What date was the SAFE agreement signed? (e.g., January 1, 2024)",
        "State of Incorporation": "In which state is the company incorporated? (e.g., Delaware)",
        "Governing Law Jurisdiction": "Which state's laws should govern this agreement?",
        "Amount": "What is the amount? (Enter amount in dollars, e.g., 1000000)",
        "Purchase Amount": "What is the purchase amount for this investment? (Enter amount in dollars, e.g., 1000000)",
        "Post-Money Valuation Cap": "What is the post-money valuation cap? (Enter amount in dollars, e.g., 10000000)",
    }

    if name in questions:
        return questions[name]

    # Default question based on type
    if placeholder["type"] == "currency":
        return f"What is the {name}? (Enter amount in dollars, e.g., 1000000)"

    return f"Please provide: {name}"


def fill_document(template_path, output_path, responses, placeholders):
    """
    Fill in placeholders in the document with provided values.

    Handles cases where placeholders are split across multiple runs.
    Handles duplicate placeholders by replacing sequentially.

    Args:
        template_path: Path to template document
        output_path: Path to save filled document
        responses: Dict mapping index to {placeholder, value}
        placeholders: List of placeholder info dicts
    """
    doc = Document(template_path)

    # Create a map of replacements to make
    # For duplicates, store list of values to use in order
    replacements_map = {}
    for idx_str in sorted(responses.keys(), key=lambda x: int(x)):
        response_data = responses[idx_str]
        placeholder_info = response_data["placeholder"]
        value = response_data["value"]
        original = placeholder_info["original"]

        # For duplicate placeholders, track by context
        if original not in replacements_map:
            replacements_map[original] = []
        replacements_map[original].append(value)

    def replace_in_paragraph(paragraph, remaining_replacements):
        """Replace placeholders in a paragraph, handling split runs and duplicates."""
        # Get full paragraph text
        full_text = paragraph.text

        # Check if any placeholder exists in this paragraph
        needs_replacement = False
        new_text = full_text

        # Replace each placeholder type, handling multiple occurrences
        for original in list(remaining_replacements.keys()):
            if original in new_text and remaining_replacements[original]:
                needs_replacement = True
                # Replace first occurrence with first remaining value
                value = remaining_replacements[original].pop(0)
                new_text = new_text.replace(original, value, 1)

                # If no more values for this placeholder, remove it
                if not remaining_replacements[original]:
                    del remaining_replacements[original]

        # If text changed, rebuild the paragraph
        if needs_replacement and new_text != full_text:
            # Clear all runs
            for run in paragraph.runs:
                run.text = ""

            # Add new text as a single run
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

    # Create a mutable copy for tracking remaining replacements
    remaining = {k: list(v) for k, v in replacements_map.items()}

    # Replace in paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para, remaining)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, remaining)

    # Replace in headers and footers
    for section in doc.sections:
        # Header
        for para in section.header.paragraphs:
            replace_in_paragraph(para, remaining)

        # Footer
        for para in section.footer.paragraphs:
            replace_in_paragraph(para, remaining)

    doc.save(output_path)


@app.route("/")
def index():
    """Render the main page."""
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload_file():
    """Handle file upload and extract placeholders."""
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]

    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not allowed_file(file.filename):
        return jsonify({"error": "Only .docx files are allowed"}), 400

    try:
        # Save uploaded file securely
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], unique_filename)
        file.save(filepath)

        # Extract placeholders
        placeholders = extract_placeholders(filepath)

        # Store in session
        session["template_path"] = filepath
        session["placeholders"] = placeholders
        session["current_index"] = 0
        session["responses"] = {}

        return jsonify(
            {
                "success": True,
                "filename": filename,
                "placeholder_count": len(placeholders),
            }
        )

    except Exception as e:
        return jsonify({"error": f"Error processing file: {str(e)}"}), 500


@app.route("/get_question", methods=["GET"])
def get_question():
    """Get the next placeholder question."""
    if "placeholders" not in session:
        return jsonify({"error": "No document uploaded"}), 400

    placeholders = session["placeholders"]
    current_index = session.get("current_index", 0)

    if current_index >= len(placeholders):
        return jsonify({"complete": True})

    placeholder = placeholders[current_index]
    question = get_placeholder_question(placeholder)

    return jsonify(
        {
            "question": question,
            "context": placeholder.get("context", ""),
            "progress": f"{current_index + 1}/{len(placeholders)}",
            "type": placeholder["type"],
        }
    )


@app.route("/submit_answer", methods=["POST"])
def submit_answer():
    """Process user's answer and move to next question."""
    if "placeholders" not in session:
        return jsonify({"error": "No document uploaded"}), 400

    data = request.json
    answer = data.get("answer", "").strip()

    if not answer:
        return jsonify({"error": "Please provide an answer"}), 400

    placeholders = session["placeholders"]
    current_index = session["current_index"]
    responses = session.get("responses", {})

    if current_index >= len(placeholders):
        return jsonify({"error": "All questions answered"}), 400

    placeholder = placeholders[current_index]

    # Format currency if needed
    if placeholder["type"] == "currency":
        try:
            # Remove common currency symbols and commas
            clean_answer = answer.replace("$", "").replace(",", "")
            amount = float(clean_answer)
            formatted_answer = f"${amount:,.0f}"
            # Store by index to handle duplicate placeholders
            responses[str(current_index)] = {
                "placeholder": placeholder,
                "value": formatted_answer,
            }
        except ValueError:
            return (
                jsonify({"error": "Please enter a valid number for currency amount"}),
                400,
            )
    else:
        # Store by index to handle duplicate placeholders
        responses[str(current_index)] = {"placeholder": placeholder, "value": answer}

    session["responses"] = responses
    session["current_index"] = current_index + 1

    return jsonify({"success": True})


@app.route("/generate", methods=["POST"])
def generate_document():
    """Generate the filled document."""
    if "template_path" not in session or "responses" not in session:
        return jsonify({"error": "Missing document or responses"}), 400

    try:
        template_path = session["template_path"]
        responses = session["responses"]
        placeholders = session["placeholders"]

        # Generate output filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"completed_{timestamp}.docx"
        output_path = os.path.join(app.config["OUTPUT_FOLDER"], output_filename)

        # Fill the document
        fill_document(template_path, output_path, responses, placeholders)

        session["output_path"] = output_path

        return jsonify({"success": True, "filename": output_filename})

    except Exception as e:
        return jsonify({"error": f"Error generating document: {str(e)}"}), 500


@app.route("/download")
def download():
    """Download the completed document."""
    if "output_path" not in session:
        return "No document available", 404

    output_path = session["output_path"]

    if not os.path.exists(output_path):
        return "Document not found", 404

    return send_file(
        output_path, as_attachment=True, download_name="completed_document.docx"
    )


@app.route("/reset", methods=["POST"])
def reset():
    """Reset the session to start over."""
    session.clear()
    return jsonify({"success": True})


if __name__ == "__main__":
    port = 5001

    # Try to create public tunnel
    try:
        from pyngrok import ngrok

        public_url = ngrok.connect(port).public_url
        print(f"\n PUBLIC URL (accessible from anywhere):")
        print(f"   {public_url}\n")
    except Exception as e:
        print(f"\n  Could not create public tunnel (network restricted)")
        print(f"   The app is running on: http://localhost:{port}")
        print(f"\n   To make it public, you can:")
        print(f"   • Use ngrok: ngrok http {port}")
        print(f"   • Use localtunnel: lt --port {port}")
        print(f"   • Deploy to: Heroku, Railway, Render, Vercel, etc.\n")

    print("📖 Open the URL above in your browser to use the app")

    # Run Flask app
    app.run(host="0.0.0.0", port=port, debug=False)
